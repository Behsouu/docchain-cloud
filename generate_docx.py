from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)


def heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
    return p


def body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '2E4057')
        shd.set(qn('w:color'), 'FFFFFF')
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)
        for run in hdr[i].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = val
    if col_widths:
        for row in table.rows:
            for c_idx, w in enumerate(col_widths):
                row.cells[c_idx].width = Inches(w)
    doc.add_paragraph()
    return table


# ── PAGE DE GARDE ──────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('EFREI Paris — Panorama du Cloud')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(52, 73, 94)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('DocChain — LegalVault')
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = RGBColor(41, 128, 185)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Plateforme de gestion securisee de documents\navec preuve d'integrite blockchain")
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(127, 140, 141)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Dossier Projet — Rapport Technique RNCP')
run.bold = True
run.font.size = Pt(12)

doc.add_paragraph()

info = [
    ('Etudiant', 'RAMEZANI Sina'),
    ('Promotion', 'EFREI Paris — Master Informatique'),
    ('Date', datetime.date.today().strftime('%d %B %Y')),
    ('URL de production', 'https://docchain-sina-efrei.azurewebsites.net'),
    ('Smart Contract Sepolia', '0xE1A55786068869B9318811f61931783600fc8FCf'),
    ('Repo GitHub', 'https://github.com/Behsouu/docchain-cloud'),
]
t = doc.add_table(rows=len(info), cols=2)
t.style = 'Table Grid'
for i, (k, v) in enumerate(info):
    t.rows[i].cells[0].text = k
    t.rows[i].cells[0].paragraphs[0].runs[0].bold = True
    t.rows[i].cells[1].text = v
    t.rows[i].cells[0].width = Inches(2.2)
    t.rows[i].cells[1].width = Inches(4.0)

doc.add_page_break()

# ── 1. INTRODUCTION ────────────────────────────────────────────────────────
heading(doc, '1. Introduction et objectifs du projet', 1, (41, 128, 185))
body(doc,
    "DocChain — LegalVault est une plateforme cloud-native developpee dans le cadre du cours "
    "Panorama du Cloud a l'EFREI Paris. Elle repond a un besoin reel de notarisation numerique : "
    "garantir l'authenticite et l'integrite de documents legaux sans tiers de confiance centralise, "
    "en combinant le stockage Azure et la blockchain Ethereum."
)

heading(doc, '1.1 Cas usage', 2)
for item in [
    "Stockage chiffre (AES-256) dans Azure Blob Storage",
    "Empreinte SHA-256 ancree sur la blockchain Ethereum Sepolia",
    "Horodatage immuable via smart contract Solidity",
    "Consultation et verification sans modifier l'original",
]:
    bullet(doc, item)

heading(doc, '1.2 Stack technique', 2)
add_table(doc,
    ['Composant', 'Technologie', 'Justification'],
    [
        ['Backend API', 'Flask 3.1 + Python 3.11', 'Leger, rapide, ideal microservices'],
        ['Hebergement', 'Azure App Service (F1 Linux)', 'PaaS manage, HTTPS natif, zip deploy'],
        ['Stockage fichiers', 'Azure Blob Storage (AES-256)', 'Scalable, prive, ecosysteme Azure'],
        ['Base de donnees', 'Azure Table Storage', 'NoSQL serverless, CRUD complet'],
        ['Secrets', 'Azure Key Vault', 'Gestion centralisee des credentials'],
        ['Reseau', 'Azure VNet + NSG', 'Isolation reseau, regles HTTPS-only'],
        ['Blockchain', 'Ethereum Sepolia + Solidity', 'Testnet, preuve immuable integrite'],
        ['IaC', 'Terraform (azurerm ~> 3.0)', 'Infrastructure reproductible, versionnee'],
        ['CI/CD', 'GitHub Actions', 'Pipeline test -> validate -> build'],
        ['Monitoring', 'Prometheus + Grafana', 'Metriques temps reel, alertes'],
    ],
    [2.0, 2.2, 2.8]
)

doc.add_page_break()

# ── 2. ARCHITECTURE ────────────────────────────────────────────────────────
heading(doc, '2. Architecture cloud et reseau', 1, (41, 128, 185))

heading(doc, '2.1 Vue ensemble infrastructure Azure', 2)
body(doc,
    "Toute l'infrastructure est deployee dans le Resource Group docchain-rg, region West Europe, "
    "via Terraform (13 ressources). Les ressources sont organisees en trois couches : "
    "reseau (VNet/NSG), compute (App Service), et donnees (Storage + Key Vault)."
)

heading(doc, '2.2 Architecture reseau — VNet et sous-reseaux', 2)
add_table(doc,
    ['Ressource', 'CIDR / Valeur', 'Role'],
    [
        ['VNet docchain-vnet', '10.0.0.0/16', 'Reseau virtuel principal (65 536 adresses)'],
        ['Subnet public', '10.0.1.0/24', 'App Service delegation Microsoft.Web/serverFarms'],
        ['Subnet prive', '10.0.2.0/24', 'Services backend : Storage, Key Vault'],
        ['NSG docchain-nsg', 'Inbound rules', 'Allow HTTPS 443 prio 100, Deny HTTP 80 prio 200'],
        ['NSG Association', 'subnet_public', 'NSG applique au subnet public'],
    ],
    [2.0, 1.8, 3.2]
)

heading(doc, '2.3 Flux de donnees complet', 2)
steps = [
    ('1. Upload', 'Client HTTPS -> App Service (gunicorn port 8000) -> calcul SHA-256'),
    ('2. Stockage', 'App Service -> Azure Blob Storage (container docchain, AES-256)'),
    ('3. Metadonnees', 'App Service -> Azure Table Storage (table documents, CRUD)'),
    ('4. Blockchain', 'Thread async -> Alchemy RPC -> Smart Contract Sepolia registerDocument'),
    ('5. Confirmation', 'TX hash -> mise a jour Table Storage (merge) -> polling frontend 5s'),
]
for name, desc in steps:
    p = doc.add_paragraph()
    run = p.add_run(name + ' : ')
    run.bold = True
    p.add_run(desc)
    p.paragraph_format.left_indent = Cm(0.5)

doc.add_paragraph()
body(doc,
    "La reponse a l'upload est immediate (< 2 secondes) car l'ancrage blockchain est asynchrone. "
    "Le frontend effectue un polling toutes les 5 secondes jusqu'a confirmation de la TX."
)

doc.add_page_break()

# ── 3. TERRAFORM ───────────────────────────────────────────────────────────
heading(doc, '3. Infrastructure as Code — Terraform', 1, (41, 128, 185))

heading(doc, '3.1 Ressources deployees (13 ressources)', 2)
add_table(doc,
    ['Ressource Terraform', 'Nom Azure', 'Type'],
    [
        ['azurerm_resource_group.rg', 'docchain-rg', 'Resource Group'],
        ['azurerm_storage_account.storage', 'docchainsinara', 'Storage Account LRS Standard'],
        ['azurerm_storage_container.container', 'docchain', 'Blob Container prive'],
        ['azurerm_storage_table.documents', 'documents', 'Table Storage CRUD metadonnees'],
        ['azurerm_service_plan.plan', 'docchain-plan', 'App Service Plan F1 Linux'],
        ['azurerm_linux_web_app.app', 'docchain-sina-efrei', 'App Service Flask Python 3.11'],
        ['azurerm_key_vault.kv', 'docchain-kv-sina', 'Key Vault standard'],
        ['azurerm_key_vault_secret.storage_key', 'storage-connection-string', 'Secret Key Vault'],
        ['azurerm_virtual_network.vnet', 'docchain-vnet', 'VNet 10.0.0.0/16'],
        ['azurerm_subnet.subnet_public', 'docchain-subnet-public', 'Subnet App Service 10.0.1.0/24'],
        ['azurerm_subnet.subnet_private', 'docchain-subnet-private', 'Subnet backend 10.0.2.0/24'],
        ['azurerm_network_security_group.nsg', 'docchain-nsg', 'NSG HTTPS allow HTTP deny'],
        ['azurerm_subnet_network_security_group_association', 'nsg_assoc', 'Association NSG subnet public'],
    ],
    [3.0, 2.0, 2.0]
)

heading(doc, '3.2 Commandes execution', 2)
for cmd, desc in [
    ('terraform init', 'Initialisation provider azurerm ~> 3.0'),
    ('terraform plan -out=tfplan', 'Generation du plan execution'),
    ('terraform apply tfplan', 'Application infrastructure'),
    ('terraform output app_url', 'URL de l app deployee'),
]:
    p = doc.add_paragraph()
    run = p.add_run(cmd)
    run.bold = True
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    p.add_run(' — ' + desc)

doc.add_page_break()

# ── 4. APPLICATION FLASK CRUD ──────────────────────────────────────────────
heading(doc, '4. Application Flask — API et CRUD complet', 1, (41, 128, 185))

heading(doc, '4.1 Endpoints REST', 2)
add_table(doc,
    ['Methode', 'Route', 'CRUD', 'Description'],
    [
        ['GET', '/', '-', 'Interface web HTML (SPA embarquee)'],
        ['POST', '/upload', 'CREATE', 'Upload fichier -> Blob -> SHA-256 -> Table -> Blockchain async'],
        ['GET', '/documents', 'READ', 'Liste tous les documents depuis Azure Table Storage'],
        ['GET', '/verify/<hash>', 'READ', 'Verifie integrite via smart contract Ethereum'],
        ['DELETE', '/documents/<row_key>', 'DELETE', 'Suppression entree Table Storage'],
        ['GET', '/metrics', '-', 'Prometheus metrics : uploads, erreurs, latence'],
        ['POST', '/alert', '-', 'Webhook Grafana alertes'],
    ],
    [1.0, 1.8, 0.8, 3.4]
)

heading(doc, '4.2 Schema Azure Table Storage', 2)
add_table(doc,
    ['Champ', 'Type', 'Description'],
    [
        ['PartitionKey', 'string', 'Valeur fixe : "documents"'],
        ['RowKey', 'UUID string', 'Identifiant unique (uuid4)'],
        ['filename', 'string', 'Nom du fichier uploade'],
        ['hash_sha256', 'string 64 hex', 'Empreinte SHA-256 du contenu'],
        ['tx_hash', 'string 66 hex', 'Hash TX Ethereum (vide si en cours)'],
        ['contract_address', 'string', 'Adresse smart contract'],
        ['file_size_bytes', 'int', 'Taille en octets'],
        ['uploaded_at', 'ISO 8601', 'Timestamp UTC upload'],
        ['status', 'string', 'Etat : "uploaded"'],
    ],
    [1.8, 1.5, 3.7]
)

heading(doc, '4.3 Upload asynchrone — performance < 2s', 2)
body(doc,
    "La cle de performance est la separation entre l'upload (synchrone < 2s) et l'ancrage "
    "blockchain (asynchrone 30-120s). threading.Thread(daemon=True) lance register_on_blockchain_async() "
    "en arriere-plan. Le frontend poll GET /documents toutes les 5s jusqu'a tx_hash confirme."
)

doc.add_page_break()

# ── 5. SMART CONTRACT ──────────────────────────────────────────────────────
heading(doc, '5. Smart Contract Ethereum — Preuve integrite', 1, (41, 128, 185))

add_table(doc,
    ['Parametre', 'Valeur'],
    [
        ['Reseau', 'Ethereum Sepolia testnet chainId 11155111'],
        ['Adresse contrat', '0xE1A55786068869B9318811f61931783600fc8FCf'],
        ['Explorateur', 'https://sepolia.etherscan.io/address/0xE1A55...'],
        ['RPC Provider', 'Alchemy eth-sepolia.g.alchemy.com/v2/'],
        ['Library Python', 'web3.py 7.15.0'],
        ['Gas limit', '200 000 (registerDocument)'],
    ],
    [2.5, 4.5]
)

add_table(doc,
    ['Fonction', 'Parametres', 'Type', 'Description'],
    [
        ['registerDocument', '_hash: bytes32, _filename: string', 'nonpayable', 'Enregistre empreinte + nom fichier'],
        ['verifyDocument', '_hash: bytes32', 'view', 'Retourne (exists, uploader, timestamp, filename)'],
        ['getDocumentCount', '—', 'view', 'Nombre total de documents'],
    ],
    [2.0, 2.5, 1.2, 2.3]
)

doc.add_page_break()

# ── 6. SECURITE ────────────────────────────────────────────────────────────
heading(doc, '6. Securite — defense en profondeur', 1, (41, 128, 185))

add_table(doc,
    ['Couche', 'Mecanisme', 'Implementation'],
    [
        ['Transport', 'TLS 1.2 minimum', 'https_only=true + NSG deny port 80'],
        ['Reseau', 'Isolation VNet', 'docchain-vnet 10.0.0.0/16 NSG HTTPS-only'],
        ['Stockage', 'Chiffrement at-rest', 'Azure Blob AES-256 gere Microsoft'],
        ['Secrets', 'Key Vault', 'storage-connection-string dans Azure Key Vault'],
        ['Integrite', 'Blockchain', 'SHA-256 + smart contract Ethereum immuable'],
        ['Acces Blob', 'Container prive', 'Aucun acces public, connection string uniquement'],
        ['Code', 'Variable sensitive', 'WALLET_PRIVATE_KEY sensitive=true Terraform'],
    ],
    [1.5, 2.0, 3.5]
)

doc.add_page_break()

# ── 7. CI/CD ───────────────────────────────────────────────────────────────
heading(doc, '7. Pipeline CI/CD — GitHub Actions', 1, (41, 128, 185))

add_table(doc,
    ['Job', 'Trigger', 'Etapes', 'Duree'],
    [
        ['test', 'push/PR main', 'checkout -> python 3.11 -> pip install (flask azure-data-tables web3) -> import test', '~45s'],
        ['terraform', 'after test', 'checkout -> terraform 1.14 -> init -backend=false -> validate', '~30s'],
        ['docker', 'after test', 'checkout -> pip install -> import test -> echo SHA', '~30s'],
    ],
    [1.2, 1.5, 3.5, 0.8]
)

body(doc,
    "Correction apportee : azure-data-tables ajoute dans le pip install du job test. "
    "Ce package est importe dans app.py (from azure.data.tables import TableServiceClient) "
    "et son absence causait un ImportError en CI."
)

doc.add_page_break()

# ── 8. DEPLOIEMENT ─────────────────────────────────────────────────────────
heading(doc, '8. Deploiement et URL de production', 1, (41, 128, 185))

heading(doc, '8.1 URL de production', 2)
p = doc.add_paragraph()
run = p.add_run('https://docchain-sina-efrei.azurewebsites.net')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(41, 128, 185)

heading(doc, '8.2 Endpoints disponibles', 2)
for route, desc in [
    ('/', 'Interface web — upload, verification, historique'),
    ('/documents', 'API JSON — liste documents (GET)'),
    ('/upload', 'API JSON — depot document (POST multipart)'),
    ('/verify/<hash>', 'API JSON — verification blockchain (GET)'),
    ('/metrics', 'Prometheus metrics (GET)'),
]:
    p = doc.add_paragraph()
    run = p.add_run(route)
    run.bold = True
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    p.add_run(' — ' + desc)
    p.paragraph_format.left_indent = Cm(0.5)

heading(doc, '8.3 Procedure de deploiement', 2)
steps_deploy = [
    ('1', 'Startup command', 'gunicorn --bind=0.0.0.0:8000 --timeout 600 --workers 2 app:app'),
    ('2', 'Package zip', 'deploy_azure/ contenant app.py + requirements.txt (sans pywin32 + gunicorn)'),
    ('3', 'az webapp deploy', '--src-path deploy_azure.zip --type zip (Kudu Oryx build ~32s)'),
    ('4', 'Verification', 'curl -> HTTP 200 en 0.33s'),
]
for num, name, desc in steps_deploy:
    p = doc.add_paragraph()
    run = p.add_run('Etape ' + num + ' — ' + name + ' : ')
    run.bold = True
    p.add_run(desc)
    p.paragraph_format.left_indent = Cm(0.3)

doc.add_page_break()

# ── 9. MONITORING ──────────────────────────────────────────────────────────
heading(doc, '9. Monitoring — Prometheus et Grafana', 1, (41, 128, 185))

add_table(doc,
    ['Metrique Prometheus', 'Type', 'Description'],
    [
        ['docchain_uploads_total', 'Counter', 'Nombre total fichiers uploades avec succes'],
        ['docchain_errors_total', 'Counter', 'Nombre total erreurs (Blob, Table, Blockchain)'],
        ['docchain_request_duration_seconds', 'Histogram', 'Duree requetes POST /upload (latence end-to-end)'],
    ],
    [3.0, 1.0, 3.0]
)

body(doc,
    "Flask expose /metrics au format Prometheus text. Grafana scrape cet endpoint via Prometheus "
    "datasource et affiche dashboard avec metriques temps reel. Alertes Grafana configurees "
    "pour notifier via webhook POST /alert."
)

doc.add_page_break()

# ── 10. BILAN ──────────────────────────────────────────────────────────────
heading(doc, '10. Bilan et competences RNCP demontrees', 1, (41, 128, 185))

add_table(doc,
    ['Competence RNCP', 'Implementation dans DocChain'],
    [
        ['Deploiement cloud PaaS', 'Azure App Service F1 Linux, gunicorn, HTTPS-only, zip deploy'],
        ['Infrastructure as Code', 'Terraform 13 ressources Azure versionnees reproductibles'],
        ['Reseau cloud securise', 'VNet 10.0.0.0/16, 2 subnets, NSG HTTPS-only, delegation Web'],
        ['Stockage cloud manage', 'Blob Storage (fichiers) + Table Storage (metadonnees CRUD)'],
        ['Secrets management', 'Azure Key Vault + variables sensitive Terraform'],
        ['CI/CD automatise', 'GitHub Actions : tests Python, terraform validate, build'],
        ['Blockchain/Web3', 'Smart contract Solidity, web3.py, ancrage SHA-256 asynchrone'],
        ['Observabilite', 'Prometheus metrics, Grafana dashboard, alertes webhook'],
        ['Securite defense profondeur', 'TLS, AES-256, container prive, NSG, Key Vault, blockchain'],
    ],
    [2.8, 4.2]
)

body(doc,
    "\nDocChain LegalVault demontre une maitrise complete du cycle de vie d'une application cloud : "
    "conception architecture, provisionnement IaC, developpement API, deploiement automatise, "
    "securisation acces et donnees, monitoring operationnel et preuve integrite blockchain. "
    "Application accessible en production depuis le 28 avril 2026."
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    'DocChain LegalVault | RAMEZANI Sina | EFREI Paris | ' +
    str(datetime.date.today().year) +
    ' | https://docchain-sina-efrei.azurewebsites.net'
)
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(127, 140, 141)

out = r'C:\EFREI\Panorama du cloud\docchain-cloud\rapport\RAMEZANI_Sina_DocChain_Dossier_Projet.docx'
doc.save(out)
print('Document cree :', out)
